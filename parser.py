"""
Resume parsing module.
Extracts text content from PDF and DOCX files.
"""
from typing import Optional, Tuple
import io
import pdfplumber
from docx import Document

class ResumeParser:
    """
    Parses resume files and extracts plain text.
    Supports PDF and DOCX formats.
    """
    
    @staticmethod
    async def parse_pdf(file_content: bytes) -> Tuple[str, dict]:
        """
        Extract text from PDF using pdfplumber.
        Returns extracted text and metadata.
        """
        try:
            text_parts = []
            metadata = {
                "pages": 0,
                "char_count": 0,
                "word_count": 0,
                "format": "pdf"
            }
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            metadata["char_count"] = len(full_text)
            metadata["word_count"] = len(full_text.split())
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"PDF parsing failed: {str(e)}")
    
    @staticmethod
    async def parse_docx(file_content: bytes) -> Tuple[str, dict]:
        """
        Extract text from DOCX using python-docx.
        Returns extracted text and metadata.
        """
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            
            metadata = {
                "pages": 1,
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "format": "docx"
            }
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"DOCX parsing failed: {str(e)}")
    
    @staticmethod
    async def parse(file_content: bytes, filename: str) -> Tuple[str, dict]:
        """
        Parse resume file based on extension.
        Routes to appropriate parser.
        """
        extension = filename.split(".")[-1].lower()
        
        if extension == "pdf":
            return await ResumeParser.parse_pdf(file_content)
        elif extension in ["docx", "doc"]:
            return await ResumeParser.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

# Convenience function
async def parse_resume(file_content: bytes, filename: str) -> Tuple[str, dict]:
    """Parse resume file and return text content."""
    return await ResumeParser.parse(file_content, filename)